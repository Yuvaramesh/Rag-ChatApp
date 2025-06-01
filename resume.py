import streamlit as st
from qdrant_client import QdrantClient
from qdrant_client.models import PointStruct, Distance, VectorParams
from sentence_transformers import SentenceTransformer
from langchain_text_splitters import RecursiveCharacterTextSplitter
from PyPDF2 import PdfReader
from docx import Document
import os
import uuid
import google.generativeai as genai

# ========== CONFIGURATION ==========
QDRANT_URL = "https://2ed85abb-e606-4167-8d2e-ce4185f33997.us-east4-0.gcp.cloud.qdrant.io"
QDRANT_API_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJhY2Nlc3MiOiJtIn0.UYr-iYmbfZzhyr-lGQBlMlMuYQIAxriQhZd6af7vLq4"
COLLECTION_NAME = "rag_chat_app"

GEMINI_API_KEY = "AIzaSyBPpxPBZbbBJdfVTowZNzHa0AOwWwxMCgk"
genai.configure(api_key=GEMINI_API_KEY)

EMBED_MODEL = SentenceTransformer("all-MiniLM-L6-v2", device="cpu")

client = QdrantClient(url=QDRANT_URL, api_key=QDRANT_API_KEY)

# ========== CREATE COLLECTION IF NOT EXISTS ==========
if not client.collection_exists(COLLECTION_NAME):
    client.create_collection(
        collection_name=COLLECTION_NAME,
        vectors_config=VectorParams(size=768, distance=Distance.COSINE)
    )

# ========== HELPERS ==========
def extract_text(file):
    ext = os.path.splitext(file.name)[-1].lower()
    if ext == '.pdf':
        reader = PdfReader(file)
        return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
    elif ext == '.docx':
        doc = Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    elif ext == '.txt':
        return file.read().decode("utf-8")
    else:
        raise ValueError(f"Unsupported file type: {ext}")

def extract_text_from_path(file_path):
    ext = os.path.splitext(file_path)[-1].lower()
    if ext == '.pdf':
        reader = PdfReader(file_path)
        return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
    elif ext == '.docx':
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    elif ext == '.txt':
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file type: {ext}")

def embed_and_store(file_name, content):
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = splitter.split_text(content)
    points = []
    for chunk in chunks:
        vector = EMBED_MODEL.encode(chunk)
        points.append(PointStruct(
            id=str(uuid.uuid4()),
            vector=vector.tolist(),
            payload={"source": file_name, "text": chunk}
        ))
    client.upsert(collection_name=COLLECTION_NAME, points=points)
    return len(chunks)

def search_context(query, top_k=3):
    query_vector = EMBED_MODEL.encode(query).tolist()
    results = client.search(
        collection_name=COLLECTION_NAME,
        query_vector=query_vector,
        limit=top_k,
        with_payload=True
    )
    return [res.payload["text"] for res in results]

def ask_gemini(prompt, context):
    full_prompt = f"""
You are an intelligent document assistant. Use the provided context to answer the user's question.

Question:
{prompt}

Context:
{context}

Answer based only on the context provided.
"""
    model = genai.GenerativeModel("gemini-1.5-flash")
    response = model.generate_content(full_prompt)
    return response.text.strip()

# ========== UI ==========
st.set_page_config(page_title="❄️ Know about Yuva Sri Ramesh", layout="wide")
st.title("👣Unveiling Yuva Sri Ramesh🎗️")

# --- Initialize session state for chat history ---
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

# --- Load Default PDF only once per session ---
if "default_doc_loaded" not in st.session_state:
    try:
        default_pdf_path = "Yuva Sri - Resume.pdf"  # Place this file in the same directory
        default_text = extract_text_from_path(default_pdf_path)
        chunk_count = embed_and_store("Yuva Sri - Resume.pdf", default_text)
        st.session_state.default_doc_loaded = True
        st.success(f"✅ Default document loaded with {chunk_count} chunks.")
    except Exception as e:
        st.error(f"Error loading default document: {e}")
        st.session_state.default_doc_loaded = False

# --- Chat Interface ---
st.divider()
st.subheader("💬 Chat with Your Documents")

query = st.text_input("Type a question or follow-up here...")

if query:
    try:
        with st.spinner("🔍 Retrieving context..."):
            context_chunks = search_context(query, top_k=3)
            context = "\n".join(context_chunks)

        with st.spinner("🤖 Thinking..."):
            answer = ask_gemini(query, context)

        # Save to chat history
        st.session_state.chat_history.append({
            "question": query,
            "answer": answer
        })

    except Exception as e:
        st.error(f"❌ Error: {e}")

# --- Display Chat History ---
if st.session_state.chat_history:
    st.markdown("### 🧠 Chat History")
    for i, chat in enumerate(st.session_state.chat_history, start=1):
        st.markdown(f"**🟣 You:** {chat['question']}")
        st.markdown(f"**🧠 Assistant:** {chat['answer']}")
        st.markdown("---")
